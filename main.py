# # import os
# # from PyPDF2 import PdfReader
# # from collections import Counter

# # from preprocessing.tokenizer import case_folding, tokenizing
# # from preprocessing.stopword import remove_stopwords
# # from preprocessing.stemming import stemming
# # from similarity.weighted_jaccard import weighted_jaccard

# # DATASET_DIR = "dataset"

# # def read_pdf(path):
# #     reader = PdfReader(path)
# #     text = ""
# #     for page in reader.pages:
# #         if page.extract_text():
# #             text += page.extract_text() + " "
# #     return text

# # def preprocess(text):
# #     text_cf = case_folding(text)
# #     tokens = tokenizing(text_cf)
# #     tokens_sw = remove_stopwords(tokens)
# #     tokens_stem = stemming(tokens_sw)
# #     return tokens, tokens_sw, tokens_stem

# # def term_frequency(tokens):
# #     return dict(Counter(tokens))

# # # ==============================
# # # INNPUT QUERY
# # # ==============================
# # query = input("\nMasukkan query pencarian: ")

# # _, _, query_tokens = preprocess(query)
# # tf_query = term_frequency(query_tokens)

# # # query_cf = case_folding(query)
# # # query_tokens = tokenizing(query_cf)
# # # query_tokens = remove_stopwords(query_tokens)
# # # query_tokens = stemming(query_tokens)
# # # tf_query = term_frequency(query_tokens)

# # print("\nTF QUERY:", tf_query)

# # # ==============================
# # # PROSES DOKUMEN
# # # ==============================
# # results = []

# # for file in os.listdir(DATASET_DIR):
# #     if file.endswith(".pdf"):
# #         print("\n==========================================================================================")
# #         print(f"FILE: {file}")
# #         print("\n==========================================================================================")

# #         text = read_pdf(os.path.join(DATASET_DIR, file))

# #         tokens, tokens_sw, tokens_stem = preprocess(text)

# #         print("Token awal:", tokens[:10])
# #         print("Setelah stopword:", tokens_sw[:10])
# #         print("Setelah stemming:", tokens_stem[:10])

# #         # tokens = stemming(
# #         #     remove_stopwords(
# #         #         tokenizing(
# #         #             case_folding(text)
# #         #         )
# #         #     )
# #         # )

# #         # Tahap 4
# #         tf = term_frequency(tokens_stem)

# #         print("\nTerm Frequency (TF):")
# #         for k, v in list(tf.items())[:10]:
# #             print(f"{k} : {v}")

# #         # tf_doc = term_frequency(tokens)

# #         score = weighted_jaccard(tf_query, tf)
# #         results.append((file, score))

# #         # # Tahap 3
# #         # text_cf = case_folding(text)
# #         # tokens = tokenizing(text_cf)
# #         # tokens_sw = remove_stopwords(tokens)
# #         # tokens_stem = stemming(tokens_sw)

# #         # print("Token awal:", tokens[:10])
# #         # print("Setelah stopword:", tokens_sw[:10])
# #         # print("Setelah stemming:", tokens_stem[:10])

# #         # # Tahap 4
# #         # tf = term_frequency(tokens_stem)

# #         # print("\nTerm Frequency (TF):")
# #         # for k, v in list(tf.items())[:10]:
# #         #     print(f"{k} : {v}")

# # # ======================
# # # HASIL RANKING
# # # ======================
# # results.sort(key=lambda x: x[1], reverse=True)

# # print("\n=== HASIL TEMU BALIK ===")
# # for file, score in results:
# #     print(f"{file} -> {score:.2f}")

# import os
# from collections import Counter, defaultdict
# from PyPDF2 import PdfReader
# from docx import Document 
# # NLP Indonesia
# from Sastrawi.Stemmer.StemmerFactory import StemmerFactory
# from Sastrawi.StopWordRemover.StopWordRemoverFactory import StopWordRemoverFactory

# from preprocessing.tokenizer import case_folding, tokenizing
# from preprocessing.stopword import remove_stopwords
# from preprocessing.stemming import stemming
# from similarity.weighted_jaccard import weighted_jaccard

# # Konfigurasi
# DATASET_DIR = "dataset"

# # Fungsi Baca File
# def read_pdf(path):
#     reader = PdfReader(path)
#     text = ""
#     for page in reader.pages:
#         if page.extract_text():
#             text += page.extract_text() + " "
#     return text

# def read_txt(path):
#     with open(path, 'r', encoding='utf-8', errors='ignore') as f:
#         return f.read()
    
# def read_docx(path):
#     doc = Document(path)
#     return " ".join([p.text for p in doc.paragraphs])

# def extract_text(path):
#     if path.endswith(".pdf"):
#         return read_pdf(path)
#     elif path.endswith(".txt"):
#         return read_txt(path)
#     elif path.endswith(".docx"):
#         return read_docx(path)
#     else:
#         return ""

# # Preprocessing
# stemmer = StemmerFactory().create_stemmer()
# stopwords = set(StopWordRemoverFactory().get_stop_words())

# def preprocess(text):
#     text_cf = case_folding(text)
#     tokens = tokenizing(text_cf)
#     tokens_sw = remove_stopwords(tokens)
#     tokens_stem = stemming(tokens_sw)
#     return tokens, tokens_sw, tokens_stem

# def term_frequency(tokens):
#     return dict(Counter(tokens))

# # Index Dokumen
# # def build_index(folder_path):
# #     index = {}

# #     for filename in os.listdir(folder_path):
# #         filepath = os.path.join(folder_path, filename)
# #         text = extract_text(filepath)

# #         if not text:
# #             continue

# #         _, _, tokens_stem = preprocess(text)
# #         tf = Counter(tokens_stem)
# #         index[filename] = tf

# #     return index
# # def build_index(folder_path):
# #     index = {}

# #     for file in os.listdir(folder_path):
# #         path = os.path.join(folder_path, file)

# #         if not file.endswith((".pdf", ".txt", ".docx")):
# #             continue

# #         text = extract_text(path)
# #         if not text:
# #             continue

# #         _, _, tokens_stem = preprocess(text)
# #         tf = Counter(tokens_stem)
# #         index[file] = tf

# #     return index

# # Inverted Index (Eksplisit)
# # def build_inverted_index(folder_path):
# #     inverted_index = defaultdict(dict)

# #     for filename in os.listdir(folder_path):
# #         filepath = os.path.join(folder_path, filename)

# #         text = extract_text(filepath)
# #         if not text:
# #             continue

# #         _, _, tokens_stem = preprocess(text)
# #         tf = Counter(tokens_stem)

# #         for term, freq in tf.items():
# #             inverted_index[term][filename] = freq

# #     return inverted_index
# # def build_inverted_index(folder_path):
# #     inverted_index = defaultdict(dict)

# #     for file in os.listdir(folder_path):
# #         path = os.path.join(folder_path, file)

# #         if not file.endswith((".pdf", ".txt", ".docx")):
# #             continue

# #         text = extract_text(path)
# #         if not text:
# #             continue

# #         _, _, tokens_stem = preprocess(text)
# #         tf = Counter(tokens_stem)

# #         for term, freq in tf.items():
# #             inverted_index[term][file] = freq

# #     return inverted_index
# # Tambahkan ini di main.py

# def build_indices(folder_path):
#     doc_index = {}
#     inverted_index = defaultdict(dict) 

#     # List file agar bisa diloop
#     files = [f for f in os.listdir(folder_path) if f.endswith((".pdf", ".txt", ".docx"))]
    
#     # (Opsional) Print info untuk debugging console
#     print(f"Memulai indexing untuk {len(files)} dokumen...")

#     for file in files:
#         path = os.path.join(folder_path, file)
        
#         # 1. Baca Teks
#         text = extract_text(path)
#         if not text: continue
            
#         # 2. Preprocessing (Ini proses paling BERAT/LAMA)
#         # Kita lakukan hanya SEKALI per file di sini
#         _, _, tokens_stem = preprocess(text)
        
#         # 3. Hitung TF
#         tf = Counter(tokens_stem)
        
#         # 4. Simpan ke Document Index
#         doc_index[file] = tf
        
#         # 5. Simpan ke Inverted Index
#         for term, freq in tf.items():
#             inverted_index[term][file] = freq

#     return doc_index, inverted_index

# # Indexing
# # def build_index(folder_path):
# #     index = {}

# #     for file in os.listdir(folder_path):
# #         path = os.path.join(folder_path, file)

# #         if file.endswith(".pdf"):
# #             text = read_pdf(path)
# #         elif file.endswith(".txt"):
# #             text = read_txt(path)
# #         elif file.endswith(".docx"):
# #             text = read_docx(path)
# #         else:
# #             continue

# #         _, _, tokens_stem = preprocess(text)
# #         tf = Counter(tokens_stem)

# #         index[file] = tf

# #     return index

# # Temu Balik & Ranking
# def search(index, tf_query):
#     scores = {}

#     for doc, tf_doc in index.items():
#         score = weighted_jaccard(tf_doc, tf_query)
#         if score > 0:
#             scores[doc] = score

#     return sorted(scores.items(), key=lambda x: x[1], reverse=True)

# # # Query
# # def process_query(query):
# #         _, _, tokens_stem = preprocess(text)
# #         tf = Counter(tokens_stem)

# # # ==============================
# # # INNPUT QUERY
# # # ==============================
# # query = input("\nMasukkan query pencarian: ")

# # _, _, query_tokens = preprocess(query)
# # tf_query = term_frequency(query_tokens)

# # print("\nTF QUERY:", tf_query)

# # # ==============================
# # # PROSES DOKUMEN
# # # ==============================
# # results = []

# # for file in os.listdir(DATASET_DIR):
# #     if file.endswith((".pdf", ".txt", ".docx")):
# #         print("\n==========================================================================================")
# #         print(f"FILE: {file}")
# #         print("\n==========================================================================================")

# #         text = extract_text(os.path.join(DATASET_DIR, file))
# #         tokens, tokens_sw, tokens_stem = preprocess(text)

# #         print("Token awal:", tokens[:10])
# #         print("Setelah stopword:", tokens_sw[:10])
# #         print("Setelah stemming:", tokens_stem[:10])

# #         # Tahap 4
# #         tf = term_frequency(tokens_stem)

# #         print("\nTerm Frequency (TF):")
# #         for k, v in list(tf.items())[:10]:
# #             print(f"{k} : {v}")

# #         score = weighted_jaccard(tf_query, tf)
# #         results.append((file, score))

# # # ======================
# # # HASIL RANKING
# # # ======================
# # results.sort(key=lambda x: x[1], reverse=True)

# # print("\n=== HASIL TEMU BALIK ===")
# # for file, score in results:
# #     print(f"{file} -> {score:.4f}")

# # # Indexing & Inverted Index
# # print("\n📚 Membuat document index...")
# # doc_index = build_index(DATASET_DIR)
# # print(f"Jumlah dokumen terindeks: {len(doc_index)}")

# # print("\n📚 Membuat inverted index...")
# # inverted_index = build_inverted_index(DATASET_DIR)
# # print(f"Jumlah term unik dalam inverted index: {len(inverted_index)}")

# # print("\n📚 CONTOH INVERTED INDEX (10 TERM PERTAMA):\n")

# # for i, (term, docs) in enumerate(inverted_index.items()):
# #     print(f"'{term}':")
# #     for doc, freq in docs.items():
# #         print(f"   {doc} -> {freq}")
# #     print()
    
# #     if i == 9:
# #         break

# # =====================================================
# # CLI MODE (DIBUNGKUS → STREAMLIT AMAN)
# # =====================================================
# if __name__ == "__main__":

#     query = input("\nMasukkan query pencarian: ")
#     _, _, query_tokens = preprocess(query)
#     tf_query = term_frequency(query_tokens)

#     print("\nTF QUERY:", tf_query)

#     results = []

#     for file in os.listdir(DATASET_DIR):
#         if file.endswith((".pdf", ".txt", ".docx")):
#             print("\n============================================")
#             print(f"FILE: {file}")
#             print("==============================================")

#             text = extract_text(os.path.join(DATASET_DIR, file))
#             tokens, tokens_sw, tokens_stem = preprocess(text)

#             print("Token awal:", tokens[:10])
#             print("Setelah stopword:", tokens_sw[:10])
#             print("Setelah stemming:", tokens_stem[:10])

#             tf = term_frequency(tokens_stem)
#             score = weighted_jaccard(tf_query, tf)
#             results.append((file, score))

#     results.sort(key=lambda x: x[1], reverse=True)

#     print("\n=== HASIL TEMU BALIK ===")
#     for file, score in results:
#         print(f"{file} -> {score:.4f}")

#     print("\n📚 Membuat Index (Optimized)...")
#     doc_index, inverted_index = build_indices(DATASET_DIR)
#     print(f"Jumlah dokumen terindeks: {len(doc_index)}")
#     print(f"Jumlah term unik: {len(inverted_index)}")

import os
import re
from collections import Counter, defaultdict
from PyPDF2 import PdfReader
from docx import Document 

# IMPORT MODUL MANUAL
from preprocessing.tokenizer import case_folding, tokenizing
from preprocessing.stopword import remove_stopwords
from preprocessing.stemming import stemming
from similarity.weighted_jaccard import weighted_jaccard

# Konfigurasi
DATASET_DIR = "dataset"

# --- FUNGSI BACA FILE (Tetap Sama) ---
def read_pdf(path):
    reader = PdfReader(path)
    text = ""
    for page in reader.pages:
        if page.extract_text():
            text += page.extract_text() + " "
    return text

def read_txt(path):
    with open(path, 'r', encoding='utf-8', errors='ignore') as f:
        return f.read()
    
def read_docx(path):
    doc = Document(path)
    return " ".join([p.text for p in doc.paragraphs])

def extract_text(path):
    if path.endswith(".pdf"):
        return read_pdf(path)
    elif path.endswith(".txt"):
        return read_txt(path)
    elif path.endswith(".docx"):
        return read_docx(path)
    else:
        return ""

# --- PREPROCESSING (MENGGUNAKAN FUNGSI MANUAL) ---
# def preprocess(text):
#     # 1. Case Folding & Tokenizing sudah digabung di tokenizer.py (rekomendasi sebelumnya)
#     # Jika tokenizer.py kamu masih terpisah, gunakan urutan ini:
#     text_cf = case_folding(text)
#     tokens = tokenizing(text_cf)
    
#     # 2. Stopword Removal Manual (Membaca file data/stopwords-id.txt)
#     tokens_sw = remove_stopwords(tokens)
    
#     # 3. Stemming Manual (Algoritma Nazief-Adriani + data/kata-dasar.original.txt)
#     tokens_stem = stemming(tokens_sw)
    
#     return tokens, tokens_sw, tokens_stem
def preprocess(text):
    # 1. Hilangkan karakter non-alfabet tapi jaga spasi
    # Ini mencegah kata terpotong paksa oleh regex [a-zA-Z]
    clean_text = re.sub(r'[^a-zA-Z\s]', ' ', text.lower())
    
    # 2. Tokenizing dengan split sederhana agar kata utuh
    tokens_orig = clean_text.split()
    
    # 3. Stopword Removal (Pastikan kata seperti 'yang', 'di' hilang)
    tokens_filt = remove_stopwords(tokens_orig)
    
    # 4. Stemming (Nazief-Adriani)
    # Tambahkan pengecekan panjang kata di sini
    tokens_stem = []
    for t in tokens_filt:
        if len(t) > 3: # Hanya stem kata yang panjangnya lebih dari 3 huruf
            stemmed = stemming([t])[0]
            tokens_stem.append(stemmed)
        else:
            tokens_stem.append(t)
            
    return tokens_orig, tokens_filt, tokens_stem

def term_frequency(tokens):
    return dict(Counter(tokens))

# --- INDEXING ---
def build_indices(folder_path):
    doc_index = {}
    inverted_index = defaultdict(dict) 

    files = [f for f in os.listdir(folder_path) if f.endswith((".pdf", ".txt", ".docx"))]
    print(f"Memulai indexing untuk {len(files)} dokumen...")

    for file in files:
        path = os.path.join(folder_path, file)
        text = extract_text(path)
        if not text: continue
            
        # Proses manual
        _, _, tokens_stem = preprocess(text)
        tf = Counter(tokens_stem)
        
        doc_index[file] = tf
        for term, freq in tf.items():
            inverted_index[term][file] = freq

    return doc_index, inverted_index

# --- MAIN EXECUTION ---
if __name__ == "__main__":
    # 1. Input Query
    query = input("\nMasukkan query pencarian: ")
    _, _, query_tokens = preprocess(query)
    tf_query = term_frequency(query_tokens)

    print("\nTF QUERY (Hasil Stemming Manual):", tf_query)

    # 2. Proses Per File (Untuk Debugging/Tampilan)
    results = []
    if not os.path.exists(DATASET_DIR):
        print(f"Folder {DATASET_DIR} tidak ditemukan!")
    else:
        for file in os.listdir(DATASET_DIR):
            if file.endswith((".pdf", ".txt", ".docx")):
                print("\n" + "="*45)
                print(f"FILE: {file}")
                print("="*45)

                text = extract_text(os.path.join(DATASET_DIR, file))
                tokens, tokens_sw, tokens_stem = preprocess(text)

                print("Token awal (10):", tokens[:10])
                print("Tanpa Stopword (10):", tokens_sw[:10])
                print("Hasil Stemming (10):", tokens_stem[:10])

                tf = term_frequency(tokens_stem)
                score = weighted_jaccard(tf_query, tf)
                results.append((file, score))

        # 3. Sorting Hasil
        results.sort(key=lambda x: x[1], reverse=True)

        print("\n" + "!"*10 + " HASIL RANKING PENCERIAN " + "!"*10)
        for i, (file, score) in enumerate(results, 1):
            print(f"{i}. {file} -> Score: {score:.4f}")

        # 4. Membuat Index (Untuk pencarian skala besar nantinya)
        print("\n📚 Membuat Inverted Index...")
        doc_index, inverted_index = build_indices(DATASET_DIR)
        print(f"Selesai! {len(doc_index)} dokumen & {len(inverted_index)} term unik terindeks.")